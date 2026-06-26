"""
build_onepager.py — generate a one-page Word briefing for a given day's
backtest, reading the figures from backtest_<tag>_SUMMARY.csv (so it works for
any date with no hard-coded numbers).

Usage:
    python build_onepager.py --tag 2026_06_19
Output: backtest_<tag>_onepager.docx
"""
import sys
from datetime import date
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1F, 0x4E, 0x78)
RED   = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x1E, 0x7A, 0x34)


def arg(flag, default):
    return sys.argv[sys.argv.index(flag) + 1] if flag in sys.argv else default


def shade(paragraph, hex_fill):
    pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_fill)
    pr.append(shd)


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)


def bullet(doc, text, lead=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1)
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.size = Pt(10)
    p.add_run(text).font.size = Pt(10)


def main():
    tag = arg("--tag", date.today().strftime("%Y_%m_%d"))
    s = pd.read_csv(f"backtest_{tag}_SUMMARY.csv").set_index("metric")
    def val(m, d=""):
        try:    return str(s.loc[m, "value"])
        except Exception: return d

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.5)
        sec.left_margin = sec.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SENSEX–NIFTY Spread Strategy — Live Run & Backtest")
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f"Trading day: {val('Date')}  |  Upstox live feed + sandbox paper-trading")
    rs.italic = True; rs.font.size = Pt(9)

    heading(doc, "The idea")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.add_run("SENSEX ≈ 3.2 × NIFTY. We track the spread = SENSEX − 3.2 × NIFTY and its z-score. "
              "When |z| ≥ 2 we bet it reverts (z ≤ −2 → buy SENSEX/sell NIFTY; z ≥ +2 → the reverse). "
              "Exit on reversion, stop-loss, or profit-target.").font.size = Pt(10)

    heading(doc, "What we built")
    bullet(doc, "pull live SENSEX & NIFTY every second from Upstox into CSV files.", "Live data → spreadsheets: ")
    bullet(doc, "each row enriched from the matching index FUTURES (volume / OI / depth).", "Full market data: ")
    bullet(doc, "places real orders on Upstox's practice account; risk limits cap losses.", "Sandbox paper-trading: ")
    bullet(doc, "applies the strategy with realistic costs + out-of-sample validation.", "Backtest engine: ")

    heading(doc, "Result (realistic, 1-minute bars, after costs)")
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
    for c, txt in zip(tbl.rows[0].cells, ("Measure", "Value", "Meaning")):
        rr = c.paragraphs[0].add_run(txt); rr.bold = True; rr.font.size = Pt(9)
    table_metrics = ["Trades taken", "Gross P&L", "Transaction cost", "Net P&L",
                     "Break-even cost", "Profitable configs",
                     "Robust profitable configs (>=5 trades)",
                     "Walk-forward (AM train / PM test)"]
    for m in table_metrics:
        if m not in s.index:
            continue
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(m).font.size = Pt(9)
        v = val(m)
        rv = cells[1].paragraphs[0].add_run(v); rv.font.size = Pt(9); rv.bold = True
        if m == "Net P&L":
            rv.font.color.rgb = GREEN if v.strip().startswith("+") else RED
        cells[2].paragraphs[0].add_run(str(s.loc[m, "what_it_means"])).font.size = Pt(9)

    heading(doc, "Bottom line")
    bl = doc.add_paragraph(); shade(bl, "FFF2CC")
    r = bl.add_run(f"{val('VERDICT')}. "); r.bold = True; r.font.size = Pt(10)
    bl.add_run(str(s.loc['VERDICT', 'what_it_means']) +
               ". One day of data is not enough to conclude anything — validate across "
               "several weeks before risking real capital.").font.size = Pt(10)

    heading(doc, "Deliverables")
    p = doc.add_paragraph()
    p.add_run(f"backtest_{tag}_report.xlsx").bold = True
    p.add_run(" (summary + every trade + all tested settings), the raw data CSVs, the paper-trade "
              "log, and reusable scripts to re-run on more data.").font.size = Pt(10)
    for rr in p.runs:
        rr.font.size = Pt(10)

    out = f"backtest_{tag}_onepager.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
